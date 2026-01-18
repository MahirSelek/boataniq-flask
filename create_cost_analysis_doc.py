"""
Script to create cost analysis document for Gemini models
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Title
title = doc.add_heading('BoataniQ - AI Model Cost Analysis', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Comprehensive Cost Estimation for Gemini 2.0 Flash and Gemini 3 Flash Models')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()  # Spacing

# Executive Summary
doc.add_heading('Executive Summary', 1)
summary = doc.add_paragraph(
    'This document provides detailed cost analysis for boat image analysis using Google Vertex AI Gemini models. '
    'Calculations are based on current pricing as of January 2025 and include three usage scenarios: '
    '1,000, 5,000, and 10,000 monthly boat image analyses.'
)
summary_format = summary.runs[0]
summary_format.font.size = Pt(11)

doc.add_paragraph()  # Spacing

# ============================================
# GEMINI 2.0 FLASH EXP SECTION
# ============================================
doc.add_heading('Gemini 2.0 Flash Exp Model Analysis', 1)

# Pricing
doc.add_heading('Pricing Structure', 2)
pricing_table = doc.add_table(rows=3, cols=2)
pricing_table.style = 'Light Grid Accent 1'
pricing_table.rows[0].cells[0].text = 'Token Type'
pricing_table.rows[0].cells[1].text = 'Price per Million Tokens'
pricing_table.rows[1].cells[0].text = 'Input Tokens'
pricing_table.rows[1].cells[1].text = '$0.10'
pricing_table.rows[2].cells[0].text = 'Output Tokens'
pricing_table.rows[2].cells[1].text = '$0.40'

# Token Usage
doc.add_heading('Token Usage per Analysis', 2)
token_table = doc.add_table(rows=4, cols=2)
token_table.style = 'Light Grid Accent 1'
token_table.rows[0].cells[0].text = 'Component'
token_table.rows[0].cells[1].text = 'Tokens'
token_table.rows[1].cells[0].text = 'Image Encoding'
token_table.rows[1].cells[1].text = '~1,290 tokens'
token_table.rows[2].cells[0].text = 'Prompt Text'
token_table.rows[2].cells[1].text = '~700 tokens'
token_table.rows[3].cells[0].text = 'Output Response'
token_table.rows[3].cells[1].text = '~2,800 tokens'

doc.add_paragraph()
total_tokens = doc.add_paragraph('Total per Analysis: 2,000 input tokens + 2,800 output tokens = 4,800 tokens')
total_tokens.runs[0].bold = True

# Cost per Analysis
doc.add_heading('Cost per Analysis', 2)
cost_calc = doc.add_paragraph()
cost_calc.add_run('Input Cost: ').bold = True
cost_calc.add_run('(2,000 / 1,000,000) × $0.10 = $0.00020\n')
cost_calc.add_run('Output Cost: ').bold = True
cost_calc.add_run('(2,800 / 1,000,000) × $0.40 = $0.00112\n')
cost_calc.add_run('Total Cost per Analysis: ').bold = True
cost_calc.add_run('$0.00132')

# Monthly Scenarios
doc.add_heading('Monthly Cost Scenarios', 2)

# Scenario 1
doc.add_heading('Scenario 1: 1,000 Monthly Analyses', 3)
scenario1_table = doc.add_table(rows=4, cols=2)
scenario1_table.style = 'Light Grid Accent 1'
scenario1_table.rows[0].cells[0].text = 'Metric'
scenario1_table.rows[0].cells[1].text = 'Value'
scenario1_table.rows[1].cells[0].text = 'Input Tokens (monthly)'
scenario1_table.rows[1].cells[1].text = '2,000,000 tokens'
scenario1_table.rows[2].cells[0].text = 'Output Tokens (monthly)'
scenario1_table.rows[2].cells[1].text = '2,800,000 tokens'
scenario1_table.rows[3].cells[0].text = 'Monthly Cost'
scenario1_table.rows[3].cells[1].text = '$1.32'
scenario1_table.rows[3].cells[0].paragraphs[0].runs[0].bold = True
scenario1_table.rows[3].cells[1].paragraphs[0].runs[0].bold = True

# Scenario 2
doc.add_heading('Scenario 2: 5,000 Monthly Analyses', 3)
scenario2_table = doc.add_table(rows=4, cols=2)
scenario2_table.style = 'Light Grid Accent 1'
scenario2_table.rows[0].cells[0].text = 'Metric'
scenario2_table.rows[0].cells[1].text = 'Value'
scenario2_table.rows[1].cells[0].text = 'Input Tokens (monthly)'
scenario2_table.rows[1].cells[1].text = '10,000,000 tokens'
scenario2_table.rows[2].cells[0].text = 'Output Tokens (monthly)'
scenario2_table.rows[2].cells[1].text = '14,000,000 tokens'
scenario2_table.rows[3].cells[0].text = 'Monthly Cost'
scenario2_table.rows[3].cells[1].text = '$6.60'
scenario2_table.rows[3].cells[0].paragraphs[0].runs[0].bold = True
scenario2_table.rows[3].cells[1].paragraphs[0].runs[0].bold = True

# Scenario 3
doc.add_heading('Scenario 3: 10,000 Monthly Analyses', 3)
scenario3_table = doc.add_table(rows=4, cols=2)
scenario3_table.style = 'Light Grid Accent 1'
scenario3_table.rows[0].cells[0].text = 'Metric'
scenario3_table.rows[0].cells[1].text = 'Value'
scenario3_table.rows[1].cells[0].text = 'Input Tokens (monthly)'
scenario3_table.rows[1].cells[1].text = '20,000,000 tokens'
scenario3_table.rows[2].cells[0].text = 'Output Tokens (monthly)'
scenario3_table.rows[2].cells[1].text = '28,000,000 tokens'
scenario3_table.rows[3].cells[0].text = 'Monthly Cost'
scenario3_table.rows[3].cells[1].text = '$13.20'
scenario3_table.rows[3].cells[0].paragraphs[0].runs[0].bold = True
scenario3_table.rows[3].cells[1].paragraphs[0].runs[0].bold = True

# Summary Table
doc.add_heading('Gemini 2.0 Flash - Summary', 2)
summary_table = doc.add_table(rows=4, cols=4)
summary_table.style = 'Light Grid Accent 1'
summary_table.rows[0].cells[0].text = 'Monthly Analyses'
summary_table.rows[0].cells[1].text = 'Input Cost'
summary_table.rows[0].cells[2].text = 'Output Cost'
summary_table.rows[0].cells[3].text = 'Total Monthly Cost'
for i in range(4):
    summary_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

summary_table.rows[1].cells[0].text = '1,000'
summary_table.rows[1].cells[1].text = '$0.20'
summary_table.rows[1].cells[2].text = '$1.12'
summary_table.rows[1].cells[3].text = '$1.32'

summary_table.rows[2].cells[0].text = '5,000'
summary_table.rows[2].cells[1].text = '$1.00'
summary_table.rows[2].cells[2].text = '$5.60'
summary_table.rows[2].cells[3].text = '$6.60'

summary_table.rows[3].cells[0].text = '10,000'
summary_table.rows[3].cells[1].text = '$2.00'
summary_table.rows[3].cells[2].text = '$11.20'
summary_table.rows[3].cells[3].text = '$13.20'

doc.add_paragraph()
annual_note = doc.add_paragraph('Annual Costs: 1,000/month = $15.84 | 5,000/month = $79.20 | 10,000/month = $158.40')
annual_note.runs[0].italic = True

doc.add_page_break()

# ============================================
# GEMINI 3 FLASH SECTION
# ============================================
doc.add_heading('Gemini 3 Flash Model Analysis', 1)

# Pricing
doc.add_heading('Pricing Structure', 2)
pricing_table3 = doc.add_table(rows=3, cols=2)
pricing_table3.style = 'Light Grid Accent 1'
pricing_table3.rows[0].cells[0].text = 'Token Type'
pricing_table3.rows[0].cells[1].text = 'Price per Million Tokens'
pricing_table3.rows[1].cells[0].text = 'Input Tokens'
pricing_table3.rows[1].cells[1].text = '$0.50'
pricing_table3.rows[2].cells[0].text = 'Output Tokens'
pricing_table3.rows[2].cells[1].text = '$3.00'

# Token Usage (same as 2.0)
doc.add_heading('Token Usage per Analysis', 2)
token_table3 = doc.add_table(rows=4, cols=2)
token_table3.style = 'Light Grid Accent 1'
token_table3.rows[0].cells[0].text = 'Component'
token_table3.rows[0].cells[1].text = 'Tokens'
token_table3.rows[1].cells[0].text = 'Image Encoding'
token_table3.rows[1].cells[1].text = '~1,290 tokens'
token_table3.rows[2].cells[0].text = 'Prompt Text'
token_table3.rows[2].cells[1].text = '~700 tokens'
token_table3.rows[3].cells[0].text = 'Output Response'
token_table3.rows[3].cells[1].text = '~2,800 tokens'

doc.add_paragraph()
total_tokens3 = doc.add_paragraph('Total per Analysis: 2,000 input tokens + 2,800 output tokens = 4,800 tokens')
total_tokens3.runs[0].bold = True

# Cost per Analysis
doc.add_heading('Cost per Analysis', 2)
cost_calc3 = doc.add_paragraph()
cost_calc3.add_run('Input Cost: ').bold = True
cost_calc3.add_run('(2,000 / 1,000,000) × $0.50 = $0.00100\n')
cost_calc3.add_run('Output Cost: ').bold = True
cost_calc3.add_run('(2,800 / 1,000,000) × $3.00 = $0.00840\n')
cost_calc3.add_run('Total Cost per Analysis: ').bold = True
cost_calc3.add_run('$0.00940')

# Monthly Scenarios
doc.add_heading('Monthly Cost Scenarios', 2)

# Scenario 1
doc.add_heading('Scenario 1: 1,000 Monthly Analyses', 3)
scenario1_table3 = doc.add_table(rows=4, cols=2)
scenario1_table3.style = 'Light Grid Accent 1'
scenario1_table3.rows[0].cells[0].text = 'Metric'
scenario1_table3.rows[0].cells[1].text = 'Value'
scenario1_table3.rows[1].cells[0].text = 'Input Tokens (monthly)'
scenario1_table3.rows[1].cells[1].text = '2,000,000 tokens'
scenario1_table3.rows[2].cells[0].text = 'Output Tokens (monthly)'
scenario1_table3.rows[2].cells[1].text = '2,800,000 tokens'
scenario1_table3.rows[3].cells[0].text = 'Monthly Cost'
scenario1_table3.rows[3].cells[1].text = '$9.40'
scenario1_table3.rows[3].cells[0].paragraphs[0].runs[0].bold = True
scenario1_table3.rows[3].cells[1].paragraphs[0].runs[0].bold = True

# Scenario 2
doc.add_heading('Scenario 2: 5,000 Monthly Analyses', 3)
scenario2_table3 = doc.add_table(rows=4, cols=2)
scenario2_table3.style = 'Light Grid Accent 1'
scenario2_table3.rows[0].cells[0].text = 'Metric'
scenario2_table3.rows[0].cells[1].text = 'Value'
scenario2_table3.rows[1].cells[0].text = 'Input Tokens (monthly)'
scenario2_table3.rows[1].cells[1].text = '10,000,000 tokens'
scenario2_table3.rows[2].cells[0].text = 'Output Tokens (monthly)'
scenario2_table3.rows[2].cells[1].text = '14,000,000 tokens'
scenario2_table3.rows[3].cells[0].text = 'Monthly Cost'
scenario2_table3.rows[3].cells[1].text = '$47.00'
scenario2_table3.rows[3].cells[0].paragraphs[0].runs[0].bold = True
scenario2_table3.rows[3].cells[1].paragraphs[0].runs[0].bold = True

# Scenario 3
doc.add_heading('Scenario 3: 10,000 Monthly Analyses', 3)
scenario3_table3 = doc.add_table(rows=4, cols=2)
scenario3_table3.style = 'Light Grid Accent 1'
scenario3_table3.rows[0].cells[0].text = 'Metric'
scenario3_table3.rows[0].cells[1].text = 'Value'
scenario3_table3.rows[1].cells[0].text = 'Input Tokens (monthly)'
scenario3_table3.rows[1].cells[1].text = '20,000,000 tokens'
scenario3_table3.rows[2].cells[0].text = 'Output Tokens (monthly)'
scenario3_table3.rows[2].cells[1].text = '28,000,000 tokens'
scenario3_table3.rows[3].cells[0].text = 'Monthly Cost'
scenario3_table3.rows[3].cells[1].text = '$94.00'
scenario3_table3.rows[3].cells[0].paragraphs[0].runs[0].bold = True
scenario3_table3.rows[3].cells[1].paragraphs[0].runs[0].bold = True

# Summary Table
doc.add_heading('Gemini 3 Flash - Summary', 2)
summary_table3 = doc.add_table(rows=4, cols=4)
summary_table3.style = 'Light Grid Accent 1'
summary_table3.rows[0].cells[0].text = 'Monthly Analyses'
summary_table3.rows[0].cells[1].text = 'Input Cost'
summary_table3.rows[0].cells[2].text = 'Output Cost'
summary_table3.rows[0].cells[3].text = 'Total Monthly Cost'
for i in range(4):
    summary_table3.rows[0].cells[i].paragraphs[0].runs[0].bold = True

summary_table3.rows[1].cells[0].text = '1,000'
summary_table3.rows[1].cells[1].text = '$1.00'
summary_table3.rows[1].cells[2].text = '$8.40'
summary_table3.rows[1].cells[3].text = '$9.40'

summary_table3.rows[2].cells[0].text = '5,000'
summary_table3.rows[2].cells[1].text = '$5.00'
summary_table3.rows[2].cells[2].text = '$42.00'
summary_table3.rows[2].cells[3].text = '$47.00'

summary_table3.rows[3].cells[0].text = '10,000'
summary_table3.rows[3].cells[1].text = '$10.00'
summary_table3.rows[3].cells[2].text = '$84.00'
summary_table3.rows[3].cells[3].text = '$94.00'

doc.add_paragraph()
annual_note3 = doc.add_paragraph('Annual Costs: 1,000/month = $112.80 | 5,000/month = $564.00 | 10,000/month = $1,128.00')
annual_note3.runs[0].italic = True

doc.add_page_break()

# ============================================
# COMPARISON SECTION
# ============================================
doc.add_heading('Model Comparison', 1)

comparison_table = doc.add_table(rows=5, cols=5)
comparison_table.style = 'Light Grid Accent 1'
comparison_table.rows[0].cells[0].text = 'Model'
comparison_table.rows[0].cells[1].text = 'Cost per Analysis'
comparison_table.rows[0].cells[2].text = '1,000/month'
comparison_table.rows[0].cells[3].text = '5,000/month'
comparison_table.rows[0].cells[4].text = '10,000/month'
for i in range(5):
    comparison_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

comparison_table.rows[1].cells[0].text = 'Gemini 2.0 Flash'
comparison_table.rows[1].cells[1].text = '$0.00132'
comparison_table.rows[1].cells[2].text = '$1.32'
comparison_table.rows[1].cells[3].text = '$6.60'
comparison_table.rows[1].cells[4].text = '$13.20'

comparison_table.rows[2].cells[0].text = 'Gemini 3 Flash'
comparison_table.rows[2].cells[1].text = '$0.00940'
comparison_table.rows[2].cells[2].text = '$9.40'
comparison_table.rows[2].cells[3].text = '$47.00'
comparison_table.rows[2].cells[4].text = '$94.00'

comparison_table.rows[3].cells[0].text = 'Difference'
comparison_table.rows[3].cells[1].text = '7.1x more'
comparison_table.rows[3].cells[2].text = '7.1x more'
comparison_table.rows[3].cells[3].text = '7.1x more'
comparison_table.rows[3].cells[4].text = '7.1x more'
for i in range(5):
    comparison_table.rows[3].cells[i].paragraphs[0].runs[0].bold = True

comparison_table.rows[4].cells[0].text = 'Annual Savings (2.0 vs 3.0)'
comparison_table.rows[4].cells[1].text = '-'
comparison_table.rows[4].cells[2].text = '$97.44'
comparison_table.rows[4].cells[3].text = '$484.80'
comparison_table.rows[4].cells[4].text = '$969.60'
for i in range(5):
    comparison_table.rows[4].cells[i].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Key Insights
doc.add_heading('Key Insights', 2)
insights = [
    '• Gemini 2.0 Flash is approximately 7.1x more cost-effective than Gemini 3 Flash',
    '• The primary cost difference comes from output token pricing ($0.40 vs $3.00 per million)',
    '• For 10,000 monthly analyses, using Gemini 2.0 Flash saves $969.60 annually',
    '• Both models use identical token counts per analysis (2,000 input + 2,800 output)',
    '• Cost scales linearly with usage volume for both models'
]

for insight in insights:
    p = doc.add_paragraph(insight, style='List Bullet')
    p.runs[0].font.size = Pt(11)

doc.add_paragraph()

# Notes
doc.add_heading('Important Notes', 2)
notes = [
    '• Pricing is based on Google Vertex AI rates as of January 2025',
    '• Token estimates assume standard boat images (1024x1024 resolution)',
    '• Output token count assumes detailed JSON analysis responses (~2,800 tokens)',
    '• Actual costs may vary based on image resolution, prompt complexity, and response length',
    '• These calculations do not include any free tier credits or promotional pricing',
    '• Monitor actual token usage in production to refine cost estimates'
]

for note in notes:
    p = doc.add_paragraph(note, style='List Bullet')
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(100, 100, 100)

# Footer
doc.add_paragraph()
footer = doc.add_paragraph('Generated: January 2025 | BoataniQ Cost Analysis')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(9)
footer.runs[0].font.color.rgb = RGBColor(150, 150, 150)

# Save document
output_path = 'BoataniQ_AI_Model_Cost_Analysis.docx'
doc.save(output_path)
print(f"✅ Document created successfully: {output_path}")
